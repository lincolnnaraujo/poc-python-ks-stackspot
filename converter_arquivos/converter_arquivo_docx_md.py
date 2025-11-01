import os
from datetime import datetime

from docx import Document


def docx_to_markdown(input_file: str, output_dir: str):
    """
    Converte um arquivo DOCX para Markdown, preservando formatação básica e imagens.

    Parâmetros:
    input_file (str): Caminho completo do arquivo DOCX
    output_dir (str): Diretório de saída para o arquivo Markdown
    Retorna:
    str: Caminho completo do arquivo Markdown gerado
    """

    try:
        # Cria o diretório de saída se não existir
        os.makedirs(output_dir, exist_ok=True)

        # Carrega o documento DOCX
        doc = Document(input_file)

        # Gera nome do arquivo de saída com timestamp
        filename = os.path.splitext(os.path.basename(input_file))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = os.path.join(output_dir, f"{filename}_{timestamp}.md")

        # Cria diretório para imagens
        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)

        markdown_lines = []
        image_counter = 1
        list_level = 0
        in_list = False

        # Processa cada parágrafo do documento
        for para in doc.paragraphs:
            text = ""
            link_url = None

            # Verifica se o parágrafo contém um hiperlink completo
            if para.hyperlinks:
                hyperlink = para.hyperlinks[0]
                text = hyperlink.text
                link_url = hyperlink.url
            else:
                # Processa runs normais
                for run in para.runs:
                    run_text = run.text.strip()
                    if not run_text:
                        continue

                    # Aplica formatação de texto
                    if run.bold and run.italic:
                        run_text = f"***{run_text}***"
                    elif run.bold:
                        run_text = f"**{run_text}**"
                    elif run.italic:
                        run_text = f"*{run_text}*"

                    text += run_text + " "

            text = text.strip()

            # Processa estrutura do documento
            if text:
                # Cabeçalhos
                if para.style.name.startswith('Heading'):
                    try:
                        level = min(6, int(para.style.name.split()[1]))
                    except (IndexError, ValueError):
                        level = 2  # Default para Heading 2
                    markdown_lines.append(f"{'#' * level} {text}")

                # Listas
                elif para.style.name == 'List Bullet':
                    if not in_list:
                        list_level = 0
                    markdown_lines.append("  " * list_level + f"- {text}")
                    in_list = True

                # Listas numeradas
                elif para.style.name == 'List Number':
                    if not in_list:
                        list_level = 0
                    markdown_lines.append("  " * list_level + f"1. {text}")
                    in_list = True

                # Blocos de citação
                elif para.style.name == 'Intense Quote':
                    markdown_lines.append(f"> {text}")

                # Linhas horizontais
                elif text.replace('-', '') == '' and len(text) >= 3:
                    markdown_lines.append("---")

                # Texto normal
                else:
                    if in_list:
                        markdown_lines.append("")
                        in_list = False
                    if link_url:
                        markdown_lines.append(f"[{text}]({link_url})")
                    else:
                        markdown_lines.append(text)
            else:
                if in_list:
                    markdown_lines.append("")
                    in_list = False

        # Processa tabelas (conversão básica)
        for table in doc.tables:
            table_md = []
            for i, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.replace('\n', '<br>'))

                if i == 0:
                    table_md.append("| " + " | ".join(row_text) + " |")
                    table_md.append("| " + " | ".join(["---"] * len(row.cells)) + " |")
                else:
                    table_md.append("| " + " | ".join(row_text) + " |")

            markdown_lines.append("\n" + "\n".join(table_md) + "\n")

        # Processa imagens
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_ext = os.path.splitext(rel.target_ref)[1].lower()
                if not image_ext:
                    image_ext = ".png"

                image_filename = f"image_{image_counter}{image_ext}"
                image_path = os.path.join(images_dir, image_filename)

                # Salva a imagem
                with open(image_path, 'wb') as f:
                    f.write(rel.target_part.blob)

                # Adiciona referência no Markdown
                markdown_lines.append(f"\n![Imagem {image_counter}](images/{image_filename})\n")
                image_counter += 1

        # Salva o arquivo Markdown
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("\n\n".join(markdown_lines))

        return "ok"
    except Exception as e:
        raise e